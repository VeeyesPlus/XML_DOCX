import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from pathlib import Path

def add_formatted_text(paragraph, element):
    """Helper function to add text with appropriate formatting to a paragraph."""
    for part in element.iter():
        if part.tag == "i":  # Italic formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.italic = True
        elif part.tag == "b":  # Bold formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.bold = True
        elif part.tag == "u":  # Underline formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.underline = True
        elif part.tag == "strike":  # Strikethrough formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.font.strike = True
        elif part.tag == "sub":  # Subscript formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.font.subscript = True
        elif part.tag == "sup":  # Superscript formatting
            run = paragraph.add_run(part.text if part.text else "")
            run.font.superscript = True
        elif part.tag == "br":  # Line break
            paragraph.add_run().add_break()
        elif part.text:
            paragraph.add_run(part.text)  # Plain text without special formatting
        if part.tail:
            paragraph.add_run(part.tail)  # Any trailing text after the tag

def parse_xml_to_docx(input_folder, output_folder):
    # Ensure output folder exists
    Path(output_folder).mkdir(parents=True, exist_ok=True)

    # Iterate over XML files in the input folder
    for xml_file in Path(input_folder).glob("*.xml"):
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # Create a new Word document for each XML file
        doc = Document()
        doc.styles['Normal'].font.size = Pt(12)

        # Find content in <text>, <references>, <work>, and <btext> tags
        for tag_name in ["text", "references", "work", "btext"]:
            for element in root.findall(f".//{tag_name}"):
                # Add formatted content to the document
                p = doc.add_paragraph()
                add_formatted_text(p, element)

        # Save the document with the same base name as the XML file in the output folder
        output_file = Path(output_folder) / f"{xml_file.stem}.docx"
        doc.save(output_file)
        print(f"Formatted content saved to {output_file}")

# Define input and output paths
input_folder = r"C:\Users\vpadimiti\Desktop\original\xmls"  # Replace with the path to your XML files
output_folder = r"C:\Users\vpadimiti\Desktop\original\docx"  # Replace with the path to save DOCX files

# Run the parsing function
parse_xml_to_docx(input_folder, output_folder)
